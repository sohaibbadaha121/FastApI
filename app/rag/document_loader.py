import os
from typing import List, Dict
from docx import Document

class DocumentLoader:
    def __init__(self, directory_path: str):
        self.directory_path = directory_path

    def load_docx_files(self) -> List[Dict[str, str]]:
        # Backward compatibility method
        return self.load_all_files()

    def load_all_files(self) -> List[Dict[str, str]]:
        documents = []
        if not os.path.exists(self.directory_path):
            return documents
            
        for filename in os.listdir(self.directory_path):
            file_path = os.path.join(self.directory_path, filename)
            if filename.endswith(".docx"):
                text = self._extract_docx_text(file_path)
                documents.append({"filename": filename, "content": text})
            elif filename.endswith(".txt"):
                text = self._extract_txt_text(file_path)
                if text.strip():
                    documents.append({"filename": filename, "content": text})
        return documents

    def _extract_docx_text(self, file_path: str) -> str:
        doc = Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells if cell.text.strip()]
                if row_data:
                    full_text.append(" | ".join(row_data))

        return "\n".join(full_text)

    def _extract_txt_text(self, file_path: str) -> str:
        # Fallbacks for Arabic encodings
        for encoding in ('utf-8', 'utf-8-sig', 'windows-1256', 'cp1256', 'latin-1'):
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception:
                continue
        return ""
